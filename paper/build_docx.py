"""
build_docx.py — Convert paper.md to paper.docx using python-docx.

Run from the repo root with the project venv activated:
  python paper/build_docx.py

Output: paper/paper.docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


HERE = Path(__file__).parent
SRC = HERE / "paper.md"
OUT = HERE / "paper.docx"


# ---------------------------------------------------------------------------
# Inline markdown -> python-docx runs
# ---------------------------------------------------------------------------

INLINE_PATTERN = re.compile(
    r"(\*\*[^*]+\*\*"          # **bold**
    r"|\*[^*]+\*"               # *italic*
    r"|`[^`]+`)"                # `code`
)


def add_runs(paragraph, text: str):
    """Add inline-formatted runs to a paragraph."""
    for piece in INLINE_PATTERN.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            run = paragraph.add_run(piece[2:-2])
            run.bold = True
        elif piece.startswith("*") and piece.endswith("*"):
            run = paragraph.add_run(piece[1:-1])
            run.italic = True
        elif piece.startswith("`") and piece.endswith("`"):
            run = paragraph.add_run(piece[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        else:
            paragraph.add_run(piece)


# ---------------------------------------------------------------------------
# Block parsers
# ---------------------------------------------------------------------------

def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    """Parse a markdown pipe-table starting at lines[start]. Returns (rows, end_idx)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        line = lines[i].strip()
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator rows (---|---)
        if all(set(c) <= {"-", ":", " "} for c in cells):
            i += 1
            continue
        rows.append(cells)
        i += 1
    return rows, i


def parse_code_block(lines: list[str], start: int) -> tuple[list[str], int]:
    """Parse a fenced code block. Returns (code_lines, end_idx after closing fence)."""
    code = []
    i = start + 1  # skip opening ```
    while i < len(lines) and not lines[i].startswith("```"):
        code.append(lines[i])
        i += 1
    return code, i + 1  # skip closing ```


# ---------------------------------------------------------------------------
# Document builders
# ---------------------------------------------------------------------------

def add_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Light Grid Accent 1"

    for r_idx, row_cells in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx in range(n_cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            text = row_cells[c_idx] if c_idx < len(row_cells) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            add_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
                if r_idx == 0:
                    run.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, code: list[str]):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("\n".join(code))
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_runs(p, item)
        p.paragraph_format.space_after = Pt(2)


# ---------------------------------------------------------------------------
# Main conversion
# ---------------------------------------------------------------------------

def convert(md_text: str) -> Document:
    doc = Document()

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Tighter heading sizes
    for name, size in [("Heading 1", 22), ("Heading 2", 16), ("Heading 3", 13)]:
        s = doc.styles[name]
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

    lines = md_text.splitlines()
    i = 0
    pending_bullets: list[str] = []

    def flush_bullets():
        nonlocal pending_bullets
        if pending_bullets:
            add_bullets(doc, pending_bullets)
            pending_bullets = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Horizontal rule -> visual separator (skip; keep clean)
        if stripped == "---":
            flush_bullets()
            i += 1
            continue

        # Blank line
        if not stripped:
            flush_bullets()
            i += 1
            continue

        # Code fence
        if stripped.startswith("```"):
            flush_bullets()
            code, i = parse_code_block(lines, i)
            add_code_block(doc, code)
            continue

        # Table
        if stripped.startswith("|"):
            flush_bullets()
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        # Headings
        if stripped.startswith("# "):
            flush_bullets()
            heading = doc.add_heading(stripped[2:].strip(), level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue
        if stripped.startswith("## "):
            flush_bullets()
            doc.add_heading(stripped[3:].strip(), level=1)
            i += 1
            continue
        if stripped.startswith("### "):
            flush_bullets()
            doc.add_heading(stripped[4:].strip(), level=2)
            i += 1
            continue
        if stripped.startswith("#### "):
            flush_bullets()
            doc.add_heading(stripped[5:].strip(), level=3)
            i += 1
            continue

        # Bullet
        m = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        if m:
            pending_bullets.append(m.group(2))
            i += 1
            continue

        # Numbered list -> just treat as paragraph with the number prefix kept
        m = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if m:
            flush_bullets()
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2))
            i += 1
            continue

        # Plain paragraph (may span multiple lines until blank)
        flush_bullets()
        para_lines = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not _is_block_start(lines[i]):
            para_lines.append(lines[i].strip())
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(para_lines))

    flush_bullets()
    return doc


def _is_block_start(line: str) -> bool:
    s = line.strip()
    return (
        s.startswith("#")
        or s.startswith("|")
        or s.startswith("```")
        or s == "---"
        or bool(re.match(r"^[-*]\s", s))
        or bool(re.match(r"^\d+\.\s", s))
    )


def main():
    md = SRC.read_text()
    doc = convert(md)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
